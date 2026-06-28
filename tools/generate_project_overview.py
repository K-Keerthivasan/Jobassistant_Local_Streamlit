from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output"
DOCX_PATH = OUTPUT_DIR / "Automatic_Resume_Generator_Project_Overview.docx"

NAVY = "17324D"
BLUE = "2563A6"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "5B6770"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_two_column_cards(doc, cards):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i in range(0, len(cards), 2):
        row = table.add_row()
        for j in range(2):
            cell = row.cells[j]
            cell.width = Inches(3.25)
            set_cell_margins(cell, 130, 160, 130, 160)
            shade(cell, LIGHT_BLUE if (i + j) % 2 == 0 else LIGHT_GRAY)
            if i + j < len(cards):
                title, body = cards[i + j]
                p = cell.paragraphs[0]
                r = p.add_run(title)
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(NAVY)
                p2 = cell.add_paragraph(body)
                p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string("24313A")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 30, NAVY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 12, BLUE),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    footer = section.footer.paragraphs[0]
    footer.add_run("Automatic Resume Generator  |  Project Overview  |  ")
    add_page_number(footer)

    # Cover
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(75)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AUTOMATIC RESUME\nGENERATOR")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Local, truth-only job-application engine")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Inches(5.9)
    set_cell_margins(cell, 240, 280, 240, 280)
    shade(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "A self-hosted platform that turns job postings into tailored, "
        "ATS-friendly resumes, cover letters, and application emails—while "
        "deterministically grounding every claim in the user’s verified profile."
    ).font.size = Pt(12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(45)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project information and technical description")
    r.italic = True
    r.font.color.rgb = RGBColor.from_string(MID_GRAY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Version 0.1.0  •  June 22, 2026").font.color.rgb = RGBColor.from_string(MID_GRAY)

    doc.add_page_break()

    # Executive summary
    add_heading(doc, "1. Executive Summary")
    doc.add_paragraph(
        "Automatic Resume Generator is a single-user, self-hosted application for "
        "managing the job-application workflow from job discovery through document "
        "generation and assisted application. Its defining constraint is “truth-only” "
        "generation: the language model may tailor wording and emphasis, but it may not "
        "invent employers, job titles, dates, education, certifications, skills, links, "
        "metrics, or other career facts."
    )
    doc.add_paragraph(
        "The system combines a browser-based Resume Studio, a Python generation pipeline, "
        "local LLM support through Ollama or Hermes, job intake and deduplication, DOCX/PDF "
        "rendering, browser-assisted form filling, and n8n-compatible automation APIs. "
        "A human review gate remains in place before any portal submission."
    )

    add_heading(doc, "2. Core Value Proposition")
    add_two_column_cards(
        doc,
        [
            ("Truth-grounded output", "A deterministic guard repairs or removes unsupported facts after model generation."),
            ("Local-first privacy", "Profile and job data remain on the user’s machine by default, with local Ollama inference."),
            ("End-to-end workflow", "Collect, review, generate, preview, export, track, and apply from one coordinated system."),
            ("ATS-friendly documents", "Structured content is rendered into reproducible DOCX and PDF application packages."),
            ("Role-aware tailoring", "Personas alter emphasis and vocabulary without changing the underlying employment history."),
            ("Human-controlled applying", "Automation fills repetitive fields and stops for review before final submission."),
        ],
    )

    add_heading(doc, "3. Main Capabilities")
    capabilities = [
        ("Resume Studio", "Web interface for job review, generation, bulk operations, scraper access, previews, and library management."),
        ("Document generation", "Creates a tailored resume, cover letter, and application email from a job description and master profile."),
        ("Truth guard", "Forces verified identity, education, links, role data, and skills; flags or strips unsupported metrics and claims."),
        ("Personas", "Automatically selects or manually applies role-specific framing such as software, full-stack, sales, or marketing."),
        ("Job intake", "Normalizes jobs from browser scraping, RSS, ATS boards, Apify, collector feeds, and manual entry."),
        ("Application assistance", "Supports n8n email workflows and visible-browser portal autofill with a final human review step."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, ("Capability", "Description")):
        shade(cell, NAVY)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
    for name, desc in capabilities:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = desc
        cells[0].paragraphs[0].runs[0].bold = True
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)

    doc.add_page_break()

    # Workflow
    add_heading(doc, "4. End-to-End Workflow")
    flow = [
        ("1", "Discover", "Jobs arrive from Tampermonkey, RSS, ATS sources, Apify, the local collector, or manual entry."),
        ("2", "Normalize & filter", "Intake applies title/location/email filters, deduplicates jobs, and places new records in the review queue."),
        ("3", "Review", "The Jobs view provides search, status, source, date, email availability, and application-state filters."),
        ("4", "Generate", "The selected profile, persona, and job description are sent to a schema-constrained local LLM."),
        ("5", "Validate", "Pydantic validates structure; the truth guard repairs identity and career facts and reports every intervention."),
        ("6", "Render", "The system produces resume and cover-letter documents plus application-email content and QA metadata."),
        ("7", "Apply", "Email jobs can route through n8n; portal jobs can be autofilled in a visible browser and stopped for user review."),
    ]
    for number, title, body in flow:
        t = doc.add_table(rows=1, cols=2)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        c1, c2 = t.rows[0].cells
        c1.width = Inches(0.55)
        c2.width = Inches(6.05)
        shade(c1, BLUE)
        shade(c2, LIGHT_GRAY)
        set_cell_margins(c1, 100, 100, 100, 100)
        set_cell_margins(c2, 100, 150, 100, 150)
        p = c1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(number)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        p = c2.paragraphs[0]
        r = p.add_run(f"{title}: ")
        r.bold = True
        p.add_run(body)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    add_heading(doc, "5. Architecture")
    doc.add_paragraph(
        "The platform uses a Python-first architecture. FastAPI serves both the HTTP API "
        "and the Resume Studio front end. The generation core builds a facts-only prompt, "
        "requests structured JSON from the configured local model, validates it, applies "
        "the truth guard, and renders documents on demand."
    )
    arch = doc.add_table(rows=5, cols=3)
    arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    arch.style = "Table Grid"
    rows = [
        ("Component", "Default port", "Responsibility"),
        ("Resume Studio API", "8088", "Web UI, generation, storage, review, and automation endpoints"),
        ("Ollama", "11434", "Local language-model inference"),
        ("Scraper collector", "8765", "Receives and displays browser-scraped job postings"),
        ("n8n", "5678", "Optional email delivery and workflow orchestration"),
    ]
    for i, row_data in enumerate(rows):
        for cell, text in zip(arch.rows[i].cells, row_data):
            cell.text = text
            set_cell_margins(cell)
            if i == 0:
                shade(cell, NAVY)
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(WHITE)
                cell.paragraphs[0].runs[0].bold = True

    add_heading(doc, "6. Truth-Only Safety Model")
    add_bullets(
        doc,
        [
            "The master profile is the only authorized source of personal and career facts.",
            "Model output is constrained to a JSON schema and validated before rendering.",
            "Identity, contact details, education, links, employers, titles, dates, and locations are restored from verified data.",
            "Skills are filtered against the profile and relevant verified skills are backfilled when necessary.",
            "Unsupported metrics, years, team sizes, and claims are flagged or removed in strict mode.",
            "A QA report records corrections and warnings for review.",
        ],
    )

    doc.add_page_break()

    # Technical detail and roadmap
    add_heading(doc, "7. Technology Stack")
    stack = [
        ("Backend", "Python 3.11+, FastAPI, Uvicorn, Pydantic"),
        ("AI engines", "Ollama and optional Hermes-compatible local gateway"),
        ("Data/config", "SQLite, YAML, JSON, environment variables"),
        ("Documents", "python-docx, LibreOffice/docx2pdf, pypdf validation"),
        ("Automation", "n8n, Playwright/Selenium, Tampermonkey"),
        ("Deployment", "Docker Compose with host-gateway connections to local services"),
        ("Front end", "Single-page Resume Studio served by the API"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Shading Accent 1"
    for area, tech in stack:
        cells = table.add_row().cells
        cells[0].text = area
        cells[1].text = tech
        cells[0].paragraphs[0].runs[0].bold = True
        for cell in cells:
            set_cell_margins(cell)

    add_heading(doc, "8. Data and Privacy")
    doc.add_paragraph(
        "The application is designed for one user per installation. Real profile files, "
        "tokens, generated state, and the SQLite database are excluded from version control. "
        "With the default Ollama configuration, profile facts and job text do not leave the "
        "user’s computer. If a hosted AI endpoint is configured, data handling becomes subject "
        "to that provider’s policies."
    )
    add_bullets(
        doc,
        [
            "Primary facts: data/profile/master_profile.yaml",
            "Role framing: data/profile/personas.yaml",
            "Application and job state: data/resume.db",
            "Autofill answers: data/apply_profile.json",
            "Job sources and filters: data/sources.yaml",
            "Secrets and service endpoints: .env",
        ],
    )

    add_heading(doc, "9. Project Structure")
    structure = [
        ("api/server.py", "FastAPI service and Resume Studio host"),
        ("src/resume_gen/", "Generation, validation, guarding, intake, rendering, and automation"),
        ("web/index.html", "Resume Studio single-page application"),
        ("data/", "Private profile/configuration templates and runtime database"),
        ("n8n/", "Importable workflow definitions and integration notes"),
        ("tampermonkey/", "Browser userscript for saving job postings"),
        ("docker/", "Container configuration"),
        ("docs/", "Setup, architecture, API, privacy, scraper, and workflow documentation"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Path", "Purpose")):
        shade(cell, NAVY)
        cell.text = text
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(WHITE)
    for path, purpose in structure:
        cells = table.add_row().cells
        cells[0].text = path
        cells[1].text = purpose
        cells[0].paragraphs[0].runs[0].font.name = "Cascadia Mono"
        for cell in cells:
            set_cell_margins(cell)

    add_heading(doc, "10. Current Status and Roadmap")
    add_bullets(
        doc,
        [
            "Generation core: implemented—resume, cover letter, email, structured validation, DOCX/PDF, CLI/API, and web UI.",
            "Intake: implemented—collector, LinkedIn/Indeed userscript, ATS/Apify sources, geographic filtering, and review queue.",
            "Auto-apply: in progress—email workflows and portal autofill operate behind review controls.",
            "Future reach: secure access from additional devices, including a planned Tailscale-based setup.",
        ],
    )

    add_heading(doc, "11. Summary")
    doc.add_paragraph(
        "Automatic Resume Generator is not merely a document writer. It is a local job-application "
        "operations platform built around verifiability, privacy, reproducibility, and user control. "
        "Its strongest architectural feature is the separation between creative language generation "
        "and deterministic fact enforcement, allowing tailored output without sacrificing factual accuracy."
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
