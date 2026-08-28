"""Render validated Resume / CoverLetter models into styled .docx files using
python-docx. Styling is defined here (acts as the template) and mirrors the
clean look of the sample resumes: centred name header, contact line, ruled
section headings, tight one-line bullets. ATS-safe (no text boxes / tables for
content, real headings, standard fonts)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from ..config import settings
from ..models import CoverLetter, Resume

ACCENT = RGBColor(0x1F, 0x4E, 0x79)  # deep blue, like the samples
DARK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x55, 0x5B, 0x66)
# Arial is deliberately boring: it is available on essentially every Windows ATS
# review workstation and has reliable metric-compatible substitutes on Linux. A
# resume should not depend on a custom font being installed to preserve its layout.
FONT = "Arial"

# Friendly labels for the master-profile skill groups, in display order.
SKILL_GROUP_LABELS = {
    "languages": "Languages",
    "backend_cloud": "Backend & Cloud",
    "frontend_ui": "Frontend",
    "ai_automation": "AI & Automation",
    "game_dev": "Game Development",
    "marketing": "Marketing",
    "sales_service": "Sales & Service",
    "creative_media": "Creative & Media",
    "office_admin": "Office & Admin",
    "tools_other": "Tools",
}


def typo(k: float = 1.0) -> dict:
    """Typography for the résumé at density ``k`` (1.0 == the original fixed
    styling, so callers that don't scale are byte-identical). Higher k → larger
    body font, looser line spacing, more section spacing, and wider margins, all
    of which make the SAME real content occupy more vertical space. The auto-fit
    search (render/autofit.py) picks the k that fills close to two pages."""
    c = lambda v, lo, hi: max(lo, min(hi, v))   # clamp
    # Font grows only gently (a 14pt résumé looks amateur); line spacing, paragraph
    # spacing and margins carry most of the vertical expansion.
    return {
        "body": round(c(11.0 * (1 + (k - 1) * 0.45), 9.5, 12.5), 2),
        "line": round(c(1.12 + (k - 1) * 0.85, 1.0, 1.5), 3),
        "after": round(c(3.0 + (k - 1) * 12, 1.5, 9), 2),
        "margin_in": round(c(0.75 + (k - 1) * 0.6, 0.6, 1.0), 3),
        "heading_size": round(c(11.5 * (1 + (k - 1) * 0.4), 10.5, 13.5), 2),
        "head_before": round(c(12.0 + (k - 1) * 14, 8, 22), 2),
        "head_after": round(c(4.0 + (k - 1) * 5, 2, 9), 2),
        "bullet_after": round(c(3.0 + (k - 1) * 10, 1.5, 9), 2),
        "bullet_line": round(c(1.1 + (k - 1) * 0.85, 1.0, 1.45), 3),
        "exp_before": round(c(7.0 + (k - 1) * 10, 4, 16), 2),
        "skills_size": round(c(10.5 * (1 + (k - 1) * 0.4), 9.5, 12.0), 2),
        "skills_after": round(c(2.0 + (k - 1) * 8, 1, 8), 2),
        "small_size": round(c(10.0 * (1 + (k - 1) * 0.4), 9.0, 11.5), 2),  # dates, years
    }


def _group_skills(skills, profile):
    """Bucket the resume's flat skill list into the master-profile categories,
    preserving the resume's relevance order within each category. Returns a list
    of (label, [skills]) for the categories that have any skills, plus 'Other'."""
    if not profile:
        return [("", list(skills))]
    groups = profile.get("skills") or {}
    # map each profile skill (lowercased) -> its group key
    skill_to_group = {}
    for gkey, items in groups.items():
        for it in items:
            skill_to_group[it.lower()] = gkey
    buckets: dict[str, list[str]] = {}
    other: list[str] = []
    for s in skills:
        g = skill_to_group.get(s.lower())
        if g:
            buckets.setdefault(g, []).append(s)
        else:
            other.append(s)
    out = []
    for gkey, label in SKILL_GROUP_LABELS.items():
        if buckets.get(gkey):
            out.append((label, buckets[gkey]))
    if other:
        out.append(("Other", other))
    return out


# --------------------------------------------------------------------------- #
# low-level helpers
# --------------------------------------------------------------------------- #
def _base_styles(doc: Document, t: dict | None = None) -> None:
    t = t or typo(1.0)
    _ats_safe_bullets(doc)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(t["body"])
    style.font.color.rgb = DARK
    pf = style.paragraph_format
    pf.space_after = Pt(t["after"])
    pf.line_spacing = t["line"]
    pf.widow_control = True  # no single orphaned line across a page break

    # Explicit page size so output is deterministic and validatable (Letter | A4).
    from docx.shared import Inches, Mm

    sec = doc.sections[0]
    if (settings.page_size or "letter").lower() == "a4":
        sec.page_width, sec.page_height = Mm(210), Mm(297)
    else:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)

    # Margins (0.75" at density 1.0): comfortable, ATS-safe; widen to fill 2 pages.
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(t["margin_in"]))


def _bottom_border(paragraph) -> None:
    """Add a thin bottom border to a paragraph (used under section headings)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1F4E79")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _run(paragraph, text, *, bold=False, size=11, color=DARK, italic=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = FONT
    # Set every Word font slot. Without this, Word/LibreOffice can substitute a
    # different font for punctuation, making PDF extraction less predictable.
    r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    return r


def _heading(doc: Document, text: str, t: dict | None = None):
    t = t or typo(1.0)
    # A built-in heading style gives DOCX parsers a semantic section boundary;
    # explicit formatting below preserves the existing visual design.
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(t["head_before"])
    p.paragraph_format.space_after = Pt(t["head_after"])
    p.paragraph_format.keep_with_next = True  # heading never sits alone at page bottom
    _run(p, text.upper(), bold=True, size=t["heading_size"], color=ACCENT)
    _bottom_border(p)
    return p


def _ats_safe_bullets(doc: Document) -> None:
    """Make list bullets extract as a real U+2022 in every renderer.

    python-docx's default template draws bullets with **F0B7 from the Symbol
    font** — a Private Use Area codepoint with no standard meaning. Word quietly
    maps it back to "•" when text is extracted, so this looks fine locally; but
    LibreOffice, which renders the PDFs served from the container, emits the raw
    PUA character. Every achievement line in the downloaded résumé then starts
    with a glyph an ATS cannot interpret.

    Rewriting the bullet levels to a literal "•" in the body font makes both
    renderers produce the same, parseable character. The list *structure* is kept,
    so DOCX parsers still see real lists.
    """
    try:
        numbering = doc.part.numbering_part.element
    except (AttributeError, KeyError, NotImplementedError, ValueError):
        return                                  # no numbering part: nothing to fix

    for lvl in numbering.iter(qn("w:lvl")):
        fmt = lvl.find(qn("w:numFmt"))
        if fmt is None or fmt.get(qn("w:val")) != "bullet":
            continue
        text_el = lvl.find(qn("w:lvlText"))
        if text_el is not None:
            text_el.set(qn("w:val"), "•")
        # The level's own run properties choose the glyph's font; point them at
        # the body font so the bullet isn't drawn from Symbol/OpenSymbol again.
        rpr = lvl.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            lvl.append(rpr)
        fonts = rpr.find(qn("w:rFonts"))
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            rpr.append(fonts)
        for slot in ("w:ascii", "w:hAnsi", "w:cs"):
            fonts.set(qn(slot), FONT)
        fonts.attrib.pop(qn("w:hint"), None)    # 'hint=default' re-selects Symbol


def _bullet(doc: Document, text: str, t: dict | None = None):
    t = t or typo(1.0)
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(t["bullet_after"])
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.line_spacing = t["bullet_line"]
    _run(p, text, size=t["body"])
    return p


def _tab_right(paragraph) -> None:
    """Configure a right-aligned tab stop at the page width for date alignment."""
    from docx.enum.text import WD_TAB_ALIGNMENT

    section = paragraph.part.document.sections[0]
    width = section.page_width - section.left_margin - section.right_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(width, WD_TAB_ALIGNMENT.RIGHT)


# --------------------------------------------------------------------------- #
# header (shared by resume + cover letter)
# --------------------------------------------------------------------------- #
def _header(doc: Document, full_name: str, contact) -> None:
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(1)
    _run(name_p, full_name.upper(), bold=True, size=22, color=ACCENT)

    parts = [contact.location, contact.email, contact.phone]
    parts += [link.url for link in contact.links]
    line = "  |  ".join(p for p in parts if p)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(4)
    _run(c, line, size=9.5, color=DARK)


# --------------------------------------------------------------------------- #
# public: resume
# ATS RULE: experience dates render as "Mon YYYY" (or "Present"). ATS parsers reliably
# read "May 2025", not "2025-05" / "05/2025" / "May 2025". Year-only values (e.g. "2023")
# are kept as-is — a month can't be invented (truth-only); fill it in the profile.
import re as _re

_MONTHS = ["", "Jan", "Feb", "Mar", "Apr", "May", "Jun",
           "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]
_MONTH_IDX = {m.lower(): i for i, m in enumerate(_MONTHS) if m}
_MONTH_IDX.update({"january": 1, "february": 2, "march": 3, "april": 4, "june": 6,
                   "july": 7, "august": 8, "september": 9, "sept": 9, "october": 10,
                   "november": 11, "december": 12})


def _norm_date(s: str) -> str:
    s = (s or "").strip()
    if not s:
        return s
    if _re.fullmatch(r"(?i)(present|current|now|ongoing|to date)", s):
        return "Present"
    m = _re.fullmatch(r"(\d{4})[-/](\d{1,2})(?:[-/]\d{1,2})?", s)            # 2024-08 / 2024/8
    if m:
        y, mo = int(m.group(1)), int(m.group(2))
        if 1 <= mo <= 12:
            return f"{_MONTHS[mo]} {y}"
    m = _re.fullmatch(r"(\d{1,2})[-/](\d{4})", s)                            # 08/2024
    if m:
        mo, y = int(m.group(1)), int(m.group(2))
        if 1 <= mo <= 12:
            return f"{_MONTHS[mo]} {y}"
    m = _re.fullmatch(r"(?i)([A-Za-z]{3,9})\.?\s+(\d{4})", s)                # August 2024 / Aug 2024
    if m and m.group(1).lower() in _MONTH_IDX:
        return f"{_MONTHS[_MONTH_IDX[m.group(1).lower()]]} {m.group(2)}"
    return s                                                                  # year-only / unknown → keep


# --------------------------------------------------------------------------- #
def render_resume(resume: Resume, out_path: Path, profile: dict | None = None,
                  density: float = 1.0) -> Path:
    t = typo(density)
    doc = Document()
    doc.core_properties.title = f"{resume.fullName} - Resume"
    doc.core_properties.subject = "Resume"
    _base_styles(doc, t)
    _header(doc, resume.fullName, resume.contact)

    if resume.summary:
        _heading(doc, "Professional Summary", t)
        p = doc.add_paragraph()
        _run(p, resume.summary, size=t["body"])

    if resume.skills:
        _heading(doc, "Skills", t)
        grouped = _group_skills(resume.skills, profile)
        if len(grouped) > 1 or (grouped and grouped[0][0]):
            # Categorized: one line per group, bold label + skills.
            for label, items in grouped:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(t["skills_after"])
                _run(p, f"{label}: ", bold=True, size=t["skills_size"], color=ACCENT)
                # Commas tokenize consistently in old and new ATS parsers.
                # Decorative glyphs can be lost during PDF extraction.
                _run(p, ", ".join(items), size=t["skills_size"])
        else:
            p = doc.add_paragraph()
            _run(p, ", ".join(resume.skills), size=t["body"])

    if resume.experience:
        _heading(doc, "Experience", t)
        for e in resume.experience:
            head = doc.add_paragraph()
            head.paragraph_format.space_before = Pt(t["exp_before"])
            head.paragraph_format.space_after = Pt(0)
            head.paragraph_format.keep_with_next = True
            _tab_right(head)
            _run(head, e.role, bold=True, size=t["body"])
            if e.company:
                _run(head, f"  |  {e.company}", size=t["body"])
            dates = f"{_norm_date(e.start)} – {_norm_date(e.end)}".strip(" –")
            _run(head, f"\t{dates}", size=t["small_size"], color=ACCENT)
            if e.location:
                loc = doc.add_paragraph()
                loc.paragraph_format.space_after = Pt(2)
                loc.paragraph_format.keep_with_next = True
                _run(loc, e.location, size=9.5, italic=True)
            for b in e.bullets:
                _bullet(doc, b, t)

    if resume.education:
        _heading(doc, "Education", t)
        for ed in resume.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            _tab_right(p)
            _run(p, ed.credential, bold=True, size=t["body"])
            _run(p, f"  |  {ed.institution}", size=t["body"])
            if ed.year:
                _run(p, f"\t{ed.year}", size=t["small_size"], color=ACCENT)

    if resume.certifications:
        _heading(doc, "Certifications", t)
        for cert in resume.certifications:
            _bullet(doc, cert, t)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


# --------------------------------------------------------------------------- #
# public: cover letter
# --------------------------------------------------------------------------- #
def render_cover_letter(cl: CoverLetter, out_path: Path) -> Path:
    doc = Document()
    _base_styles(doc)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(1)
    _run(name_p, cl.fullName.upper(), bold=True, size=20, color=ACCENT)

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    _run(c, cl.contactLine, size=9.5)

    g = doc.add_paragraph()
    g.paragraph_format.space_after = Pt(6)
    _run(g, cl.greeting)

    for para in cl.body:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        _run(p, para)

    s = doc.add_paragraph()
    s.paragraph_format.space_before = Pt(6)
    _run(s, cl.signOff)
    sig = doc.add_paragraph()
    _run(sig, cl.signature)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
