from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document

from resume_gen.models import Contact, EducationItem, ExperienceItem, Link, Resume
from resume_gen.render.docx_renderer import render_resume


class ResumeDocxAtsTests(unittest.TestCase):
    def _resume(self) -> Resume:
        return Resume(
            fullName="Jordan Sample",
            headline="Technology Adoption Consultant",
            contact=Contact(
                email="jordan@example.com",
                phone="555-0100",
                location="Toronto, ON",
                links=[Link(label="LinkedIn", url="https://linkedin.com/in/jordan")],
            ),
            summary="Consultant experienced in API integration and client delivery.",
            skills=["Python", "REST APIs", "Client relationship management"],
            experience=[
                ExperienceItem(
                    company="Example Co",
                    role="Consultant",
                    location="Toronto, ON",
                    start="2024-01",
                    end="present",
                    bullets=["Built API integrations for client workflows."],
                )
            ],
            education=[
                EducationItem(
                    institution="Example College",
                    credential="Diploma in Technology",
                    year="2023",
                )
            ],
        )

    def test_resume_uses_plain_semantic_ats_structure(self):
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            render_resume(self._resume(), path)
            doc = Document(path)

        texts = [paragraph.text for paragraph in doc.paragraphs]
        self.assertEqual(texts[0], "JORDAN SAMPLE")
        # The headline/title is deliberately NOT rendered: the name is followed
        # straight by the contact line. `Resume.headline` still exists on the
        # model, it just doesn't reach the page.
        self.assertIn("jordan@example.com", texts[1])
        self.assertNotIn("Technology Adoption Consultant", "\n".join(texts))
        self.assertEqual(len(doc.tables), 0)
        self.assertNotIn("•", "\n".join(texts[:8]))

        heading_texts = [
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.style.name == "Heading 2"
        ]
        self.assertEqual(
            heading_texts,
            ["PROFESSIONAL SUMMARY", "SKILLS", "EXPERIENCE", "EDUCATION"],
        )

        extracted = "\n".join(texts)
        for expected in (
            "jordan@example.com",
            "Python, REST APIs, Client relationship management",
            "Jan 2024",
            "Present",
        ):
            self.assertIn(expected, extracted)


    def test_bullets_are_a_real_unicode_bullet_not_a_symbol_glyph(self):
        """Guards the ATS bug where every achievement line began with U+F0B7.

        python-docx's default List Bullet numbering draws bullets with F0B7 from
        the Symbol font — a Private Use Area codepoint. Word maps it back to "•"
        when extracting text, so it looks fine locally, but LibreOffice (which
        renders the PDFs served from the container) emits the raw PUA character
        and an ATS then reads an unparseable glyph before every bullet.
        """
        import re
        import zipfile

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            render_resume(self._resume(), path)
            with zipfile.ZipFile(path) as z:
                numbering = z.read("word/numbering.xml").decode("utf-8")
                document = z.read("word/document.xml").decode("utf-8")

        self.assertNotIn("", numbering, "Symbol-font bullet is back")
        self.assertNotIn("", document)
        bullet_levels = [
            m for m in re.findall(r'w:lvlText w:val="([^"]*)"', numbering)
            if m and not m.startswith("%")
        ]
        self.assertTrue(bullet_levels, "no bullet levels found in numbering.xml")
        for glyph in bullet_levels:
            self.assertEqual(glyph, "•", f"bullet glyph is {glyph!r}, not U+2022")

    def test_no_layout_constructs_that_ats_parsers_choke_on(self):
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "resume.docx"
            render_resume(self._resume(), path)
            import zipfile

            with zipfile.ZipFile(path) as z:
                document = z.read("word/document.xml").decode("utf-8")

        for tag, label in (("<w:tbl>", "tables"), ("<w:drawing>", "images"),
                           ("<w:txbxContent>", "text boxes")):
            self.assertNotIn(tag, document, f"resume contains {label}")


if __name__ == "__main__":
    unittest.main()
